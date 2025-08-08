#!/usr/bin/env python3
"""
Resume file processing utilities for extracting text from uploaded documents
Supports PDF, DOC, and DOCX file formats
"""

import os
import re
import io
import logging
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import tempfile
import uuid

# File processing libraries
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ResumeFileProcessor:
    """Handles file upload and text extraction from resume documents"""
    
    SUPPORTED_FORMATS = ['.pdf', '.doc', '.docx']
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        self.upload_dir = os.path.join(os.path.dirname(__file__), 'uploads')
        self.ensure_upload_directory()
    
    def ensure_upload_directory(self):
        """Ensure upload directory exists"""
        if not os.path.exists(self.upload_dir):
            os.makedirs(self.upload_dir, exist_ok=True)
            logger.info(f"Created upload directory: {self.upload_dir}")
    
    def validate_file(self, file, filename: str) -> Tuple[bool, str]:
        """Validate uploaded file format and size"""
        try:
            # Check file extension
            if not filename:
                return False, "No filename provided"
            
            file_ext = os.path.splitext(filename.lower())[1]
            if file_ext not in self.SUPPORTED_FORMATS:
                return False, f"Unsupported file format. Supported formats: {', '.join(self.SUPPORTED_FORMATS)}"
            
            # Check file size
            file.seek(0, 2)  # Seek to end
            file_size = file.tell()
            file.seek(0)  # Reset to beginning
            
            if file_size > self.MAX_FILE_SIZE:
                return False, f"File size exceeds maximum limit of {self.MAX_FILE_SIZE // (1024 * 1024)}MB"
            
            if file_size == 0:
                return False, "File is empty"
            
            return True, ""
            
        except Exception as e:
            logger.error(f"Error validating file: {e}")
            return False, f"Error validating file: {str(e)}"
    
    def save_uploaded_file(self, file, filename: str) -> Tuple[bool, str, str]:
        """Save uploaded file to temporary location"""
        try:
            # Generate unique filename
            file_ext = os.path.splitext(filename.lower())[1]
            unique_filename = f"{uuid.uuid4()}{file_ext}"
            file_path = os.path.join(self.upload_dir, unique_filename)
            
            # Save file
            file.seek(0)
            with open(file_path, 'wb') as f:
                f.write(file.read())
            
            logger.info(f"File saved: {file_path}")
            return True, file_path, unique_filename
            
        except Exception as e:
            logger.error(f"Error saving file: {e}")
            return False, "", f"Error saving file: {str(e)}"
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[bool, str, str]:
        """Extract text from PDF file"""
        try:
            text = ""
            
            # Try pdfplumber first (better for complex layouts)
            if PDFPLUMBER_AVAILABLE:
                try:
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                    
                    if text.strip():
                        return True, text, ""
                except Exception as e:
                    logger.warning(f"pdfplumber extraction failed: {e}")
            
            # Fallback to PyPDF2
            if PDF_AVAILABLE:
                try:
                    with open(file_path, 'rb') as f:
                        pdf_reader = PdfReader(f)
                        for page in pdf_reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                    
                    if text.strip():
                        return True, text, ""
                except Exception as e:
                    logger.warning(f"PyPDF2 extraction failed: {e}")
            
            return False, "", "Unable to extract text from PDF. The file may be scanned or corrupted."
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return False, "", f"Error processing PDF: {str(e)}"
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[bool, str, str]:
        """Extract text from DOCX file"""
        try:
            text = ""
            
            # Try python-docx first
            if DOCX_AVAILABLE:
                try:
                    doc = Document(file_path)
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text += cell.text + "\n"
                    
                    if text.strip():
                        return True, text, ""
                except Exception as e:
                    logger.warning(f"python-docx extraction failed: {e}")
            
            # Try mammoth as fallback
            if MAMMOTH_AVAILABLE:
                try:
                    with open(file_path, "rb") as docx_file:
                        result = mammoth.extract_raw_text(docx_file)
                        text = result.value
                    
                    if text.strip():
                        return True, text, ""
                except Exception as e:
                    logger.warning(f"mammoth extraction failed: {e}")
            
            return False, "", "Unable to extract text from DOCX file. Please ensure the file is not corrupted."
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return False, "", f"Error processing DOCX: {str(e)}"
    
    def extract_text_from_doc(self, file_path: str) -> Tuple[bool, str, str]:
        """Extract text from DOC file"""
        try:
            # Try mammoth (best for .doc files)
            if MAMMOTH_AVAILABLE:
                try:
                    with open(file_path, "rb") as doc_file:
                        result = mammoth.extract_raw_text(doc_file)
                        text = result.value
                    
                    if text.strip():
                        return True, text, ""
                except Exception as e:
                    logger.warning(f"mammoth extraction failed: {e}")
            
            return False, "", "Unable to extract text from DOC file. Please convert to DOCX format for better compatibility."
            
        except Exception as e:
            logger.error(f"Error extracting text from DOC: {e}")
            return False, "", f"Error processing DOC: {str(e)}"
    
    def extract_text_from_file(self, file_path: str) -> Tuple[bool, str, str]:
        """Extract text from any supported file format"""
        try:
            file_ext = os.path.splitext(file_path.lower())[1]
            
            if file_ext == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_ext == '.docx':
                return self.extract_text_from_docx(file_path)
            elif file_ext == '.doc':
                return self.extract_text_from_doc(file_path)
            else:
                return False, "", f"Unsupported file format: {file_ext}"
                
        except Exception as e:
            logger.error(f"Error extracting text from file: {e}")
            return False, "", f"Error processing file: {str(e)}"
    
    def cleanup_file(self, file_path: str):
        """Clean up temporary file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Cleaned up file: {file_path}")
        except Exception as e:
            logger.warning(f"Error cleaning up file {file_path}: {e}")

class ResumeTextParser:
    """Parses extracted resume text and identifies sections"""
    
    def __init__(self):
        self.section_patterns = {
            'personal_info': [
                r'contact\s+information',
                r'personal\s+details',
                r'contact\s+details'
            ],
            'summary': [
                r'professional\s+summary',
                r'career\s+summary',
                r'summary',
                r'profile',
                r'objective',
                r'career\s+objective'
            ],
            'experience': [
                r'work\s+experience',
                r'professional\s+experience',
                r'employment\s+history',
                r'experience',
                r'career\s+history'
            ],
            'education': [
                r'education',
                r'educational\s+background',
                r'academic\s+background'
            ],
            'skills': [
                r'skills',
                r'technical\s+skills',
                r'core\s+competencies',
                r'competencies',
                r'areas\s+of\s+expertise'
            ],
            'certifications': [
                r'certifications',
                r'licenses',
                r'certificates'
            ]
        }
    
    def parse_resume_text(self, text: str) -> Dict[str, Any]:
        """Parse resume text and extract structured data"""
        try:
            # Clean and normalize text
            cleaned_text = self.clean_text(text)
            
            # Extract sections
            sections = self.identify_sections(cleaned_text)
            
            # Parse each section
            parsed_data = {
                'personal_info': self.parse_personal_info(sections.get('personal_info', '') or cleaned_text[:500]),
                'summary': self.parse_summary(sections.get('summary', '')),
                'work_experience': self.parse_work_experience(sections.get('experience', '')),
                'education': self.parse_education(sections.get('education', '')),
                'skills': self.parse_skills(sections.get('skills', '')),
                'certifications': self.parse_certifications(sections.get('certifications', ''))
            }
            
            # Generate extraction summary
            summary = self.generate_extraction_summary(parsed_data)
            
            return {
                'success': True,
                'extracted_data': parsed_data,
                'extraction_summary': summary,
                'raw_text': text
            }
            
        except Exception as e:
            logger.error(f"Error parsing resume text: {e}")
            return {
                'success': False,
                'error': f"Error parsing resume: {str(e)}",
                'raw_text': text
            }
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere
        text = re.sub(r'[^\w\s@.-]', ' ', text)
        
        # Normalize line breaks
        text = re.sub(r'\s*\n\s*', '\n', text)
        
        return text.strip()
    
    def identify_sections(self, text: str) -> Dict[str, str]:
        """Identify and extract different sections from resume text"""
        sections = {}
        text_lower = text.lower()
        
        for section_name, patterns in self.section_patterns.items():
            for pattern in patterns:
                matches = list(re.finditer(pattern, text_lower))
                if matches:
                    # Find the start of this section
                    start_pos = matches[0].start()
                    
                    # Find the end (start of next section or end of text)
                    end_pos = len(text)
                    for other_section, other_patterns in self.section_patterns.items():
                        if other_section != section_name:
                            for other_pattern in other_patterns:
                                other_matches = list(re.finditer(other_pattern, text_lower[start_pos + 50:]))
                                if other_matches:
                                    potential_end = start_pos + 50 + other_matches[0].start()
                                    if potential_end < end_pos:
                                        end_pos = potential_end
                    
                    sections[section_name] = text[start_pos:end_pos].strip()
                    break
        
        return sections
    
    def parse_personal_info(self, text: str) -> Dict[str, str]:
        """Extract personal information"""
        info = {}
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            info['email'] = email_match.group()
        
        # Phone
        phone_patterns = [
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
            r'\(\d{3}\)\s*\d{3}[-.]?\d{4}',
            r'\b\d{10}\b'
        ]
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                info['phone'] = phone_match.group()
                break
        
        # Name (assume first few words before email/phone)
        lines = text.split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if line and len(line.split()) <= 4 and not re.search(r'[@\d]', line):
                info['name'] = line
                break
        
        # Location (look for state abbreviations or city patterns)
        location_pattern = r'\b[A-Z][a-z]+,?\s+[A-Z]{2}\b|\b[A-Z][a-z]+\s+[A-Z][a-z]+,?\s+[A-Z]{2}\b'
        location_match = re.search(location_pattern, text)
        if location_match:
            info['location'] = location_match.group()
        
        return info
    
    def parse_summary(self, text: str) -> str:
        """Extract professional summary"""
        if not text:
            return ""
        
        # Remove section header
        lines = text.split('\n')
        summary_lines = []
        
        for line in lines:
            line = line.strip()
            if line and not any(header in line.lower() for header in ['summary', 'profile', 'objective']):
                summary_lines.append(line)
        
        return ' '.join(summary_lines)
    
    def parse_work_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience entries"""
        if not text:
            return []
        
        experiences = []
        
        # Split by potential job entries (look for patterns like dates, companies)
        # This is a simplified approach - more sophisticated parsing could be added
        lines = text.split('\n')
        current_exp = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Date pattern (various formats)
            date_pattern = r'\b\d{1,2}\/\d{4}\b|\b\d{4}\b|\b[A-Z][a-z]+\s+\d{4}\b'
            if re.search(date_pattern, line):
                if current_exp:
                    experiences.append(current_exp)
                current_exp = {'dates': line}
            elif 'title' not in current_exp and len(line.split()) <= 6:
                current_exp['title'] = line
            elif 'company' not in current_exp and len(line.split()) <= 4:
                current_exp['company'] = line
            else:
                if 'description' not in current_exp:
                    current_exp['description'] = line
                else:
                    current_exp['description'] += ' ' + line
        
        if current_exp:
            experiences.append(current_exp)
        
        return experiences[:5]  # Limit to 5 entries
    
    def parse_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education entries"""
        if not text:
            return []
        
        education = []
        lines = text.split('\n')
        
        current_edu = {}
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Look for degree patterns
            degree_keywords = ['bachelor', 'master', 'associate', 'diploma', 'certificate', 'phd', 'doctorate']
            if any(keyword in line.lower() for keyword in degree_keywords):
                if current_edu:
                    education.append(current_edu)
                current_edu = {'degree': line}
            elif 'institution' not in current_edu:
                current_edu['institution'] = line
            elif re.search(r'\b\d{4}\b', line):
                current_edu['graduation_year'] = re.search(r'\b\d{4}\b', line).group()
        
        if current_edu:
            education.append(current_edu)
        
        return education[:3]  # Limit to 3 entries
    
    def parse_skills(self, text: str) -> Dict[str, List[str]]:
        """Extract skills"""
        if not text:
            return {'technical': [], 'soft': []}
        
        # Common technical keywords
        technical_keywords = [
            'microsoft', 'excel', 'word', 'powerpoint', 'office', 'computer',
            'software', 'database', 'programming', 'html', 'css', 'javascript',
            'python', 'java', 'sql', 'windows', 'mac', 'linux'
        ]
        
        # Common soft skill keywords
        soft_keywords = [
            'communication', 'leadership', 'teamwork', 'management', 'organization',
            'planning', 'problem', 'analytical', 'creative', 'customer', 'service'
        ]
        
        skills = {'technical': [], 'soft': []}
        
        # Split by common delimiters
        skill_items = re.split(r'[,;\n•·]', text)
        
        for item in skill_items:
            item = item.strip()
            if not item or len(item) < 3:
                continue
            
            # Categorize as technical or soft skill
            item_lower = item.lower()
            if any(keyword in item_lower for keyword in technical_keywords):
                if item not in skills['technical']:
                    skills['technical'].append(item)
            elif any(keyword in item_lower for keyword in soft_keywords):
                if item not in skills['soft']:
                    skills['soft'].append(item)
            else:
                # Default to technical if unclear
                if item not in skills['technical'] and len(skills['technical']) < 10:
                    skills['technical'].append(item)
        
        return skills
    
    def parse_certifications(self, text: str) -> List[Dict[str, str]]:
        """Extract certifications"""
        if not text:
            return []
        
        certifications = []
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if line and len(line) > 5:  # Minimum reasonable certification name length
                cert = {'name': line}
                
                # Look for year in the line
                year_match = re.search(r'\b\d{4}\b', line)
                if year_match:
                    cert['date_obtained'] = year_match.group()
                
                certifications.append(cert)
        
        return certifications[:5]  # Limit to 5 entries
    
    def generate_extraction_summary(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate summary of extracted data"""
        summary = {
            'sections_found': 0,
            'experience_entries': len(parsed_data.get('work_experience', [])),
            'education_entries': len(parsed_data.get('education', [])),
            'skills_count': (
                len(parsed_data.get('skills', {}).get('technical', [])) + 
                len(parsed_data.get('skills', {}).get('soft', []))
            ),
            'certifications_count': len(parsed_data.get('certifications', [])),
            'has_personal_info': bool(parsed_data.get('personal_info', {})),
            'has_summary': bool(parsed_data.get('summary', ''))
        }
        
        # Count sections with content
        if summary['has_personal_info']:
            summary['sections_found'] += 1
        if summary['has_summary']:
            summary['sections_found'] += 1
        if summary['experience_entries'] > 0:
            summary['sections_found'] += 1
        if summary['education_entries'] > 0:
            summary['sections_found'] += 1
        if summary['skills_count'] > 0:
            summary['sections_found'] += 1
        if summary['certifications_count'] > 0:
            summary['sections_found'] += 1
        
        return summary

# Example usage
if __name__ == "__main__":
    processor = ResumeFileProcessor()
    parser = ResumeTextParser()
    
    print("Resume File Processor Test")
    print(f"Supported formats: {processor.SUPPORTED_FORMATS}")
    print(f"PDF extraction available: {PDF_AVAILABLE or PDFPLUMBER_AVAILABLE}")
    print(f"DOCX extraction available: {DOCX_AVAILABLE}")
    print(f"DOC extraction available: {MAMMOTH_AVAILABLE}")
    print("✅ Resume file processor initialized successfully!")